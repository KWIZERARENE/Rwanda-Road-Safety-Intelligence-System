"""
Technical Report Docx Generator for RRSIS Project
--------------------------------------------------
Converts Technical_Report.md or builds formatted docs/Technical_Report.docx
using python-docx.

Run script:
    python scripts/generate_docx_report.py
"""

import os
import sys

def create_docx_report():
    try:
        import docx
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        print("[NOTICE] 'python-docx' library is not installed.")
        print("To generate 'Technical_Report.docx', run: pip install python-docx")
        return False

    doc = Document()

    # Set standard page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("RWANDA ROAD SAFETY INTELLIGENCE SYSTEM (RRSIS)\n")
    run_title.bold = True
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = RGBColor(0, 51, 102)

    run_sub = p_title.add_run("HDFS + Apache Spark (PySpark DataFrame) Analytics Engine\n")
    run_sub.bold = True
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = RGBColor(100, 100, 100)

    run_meta = p_title.add_run("Mid-Term Group Project Technical Report & Methodology\n")
    run_meta.italic = True
    run_meta.font.size = Pt(11)

    doc.add_paragraph("=" * 60)

    # Reading Markdown report content
    md_path = os.path.join("docs", "Technical_Report.md")
    if not os.path.exists(md_path):
        md_path = os.path.join("..", "docs", "Technical_Report.md")

    if os.path.exists(md_path):
        with open(md_path, "r", encoding="utf-8") as f:
            lines = f.readlines()

        for line in lines:
            line_str = line.strip()
            if line_str.startswith("# "):
                p = doc.add_heading(level=1)
                run = p.add_run(line_str.replace("# ", ""))
                run.font.color.rgb = RGBColor(0, 51, 102)
            elif line_str.startswith("## "):
                p = doc.add_heading(level=2)
                run = p.add_run(line_str.replace("## ", ""))
                run.font.color.rgb = RGBColor(0, 102, 153)
            elif line_str.startswith("### "):
                p = doc.add_heading(level=3)
                run = p.add_run(line_str.replace("### ", ""))
                run.font.color.rgb = RGBColor(51, 51, 51)
            elif line_str.startswith("- ") or line_str.startswith("* "):
                doc.add_paragraph(line_str[2:], style='List Bullet')
            elif line_str == "---" or line_str.startswith("==="):
                p = doc.add_paragraph()
                p.add_run("-" * 50)
            elif line_str:
                doc.add_paragraph(line_str)
    else:
        doc.add_paragraph("Technical Report Content placeholder.")

    # Save to docs/Technical_Report.docx
    output_dir = "docs"
    os.makedirs(output_dir, exist_ok=True)
    out_path = os.path.join(output_dir, "Technical_Report.docx")
    doc.save(out_path)
    print(f"[SUCCESS] Technical Report saved to: {out_path}")
    return True


if __name__ == "__main__":
    create_docx_report()
